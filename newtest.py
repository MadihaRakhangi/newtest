import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tabulate import tabulate
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import csv

F = pd.read_csv("main.csv")


def find_loc_name(loc_id, locations):
    while loc_id in locations:
        loc_data = locations[loc_id]
        loc_name = loc_data["loc_name"]
        granularity = int(loc_data["granularity"])
        loc_type = loc_data["type"]

        if granularity == 1:
            if loc_type != "fac_area":
                parent_id = loc_data["parent_id"]
                parent_loc_name = find_loc_name(parent_id, locations)
                return parent_loc_name if parent_loc_name else "Parent Location"
            else:
                return loc_name
        else:
            loc_id = loc_data["parent_id"]
    return None


def create_dataframe(F):
    with open("your_file.csv", "r") as file:
        reader = csv.DictReader(file)
        locations = {row["loc_id"]: row for row in reader}
    table_data = []
    for index, row in F.iterrows():
        loc_id = str(row["loc_id"])
        result = find_loc_name(loc_id, locations)
        if result:
            loc_name = locations[loc_id]["loc_name"]
            parent_id = locations[loc_id]["parent_id"]
            parent_loc_name = locations[parent_id]["loc_name"]
            table_data.append([loc_name, parent_loc_name if parent_loc_name else "", result])
    df = pd.DataFrame(table_data, columns=["Location", "Parent Location", "Facility Area"])
    return df


table_df = create_dataframe(F)

phs = pd.DataFrame()  # OP
phs["VL1-L2 (V)"] = F["op_l1_l2_v"]
phs["VL2-L3 (V)"] = F["op_l2_l3_v"]
phs["VL3-L1 (V)"] = F["op_l3_l1_v"]
phs["VL1-N (V)"] = F["op_l1_n_v"]
phs["VL2-N (V)"] = F["op_l2_n_v"]
phs["VL3-N (V)"] = F["op_l3_n_v"]
phs["Phase Sequence"] = F["phase_seq"]


def phase_result(phase_seq):
    if phase_seq == "RYB":
        return "CLOCKWISE"
    elif phase_seq == "RBY":
        return "ANTICLOCKWISE"
    else:
        return ""


phs.loc[:, "Result"] = phs["Phase Sequence"].apply(phase_result)


def phs_table(df, doc):
    table_data = df.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False

    max_word_lengths = [0] * num_cols
    for j, col in enumerate(table_data.columns):
        max_word_length = max(table_data[col].astype(str).apply(len))
        max_word_lengths[j] = max_word_length
    conversion_factor = 0.1
    column_widths = [length * conversion_factor for length in max_word_lengths]

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=0):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i + 1, j).text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 8
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def main():
    phs_df = pd.concat([table_df, phs], axis=1)
    phs_df = phs_df.dropna()
    print(phs_df)
    doc = Document()
    doc.add_heading("Phase Sequence test", 0)
    doc = phs_table(phs_df, doc)
    doc.save("newtest.docx")


main()
