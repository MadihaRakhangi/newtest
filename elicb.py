import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
import matplotlib.pyplot as plt
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

fg = pd.read_csv("sugg-max-eli.csv")
F = pd.read_csv("main.csv")

elicbdf = pd.DataFrame()
elicbdf["Earthing Configuration"] = F["earth_config"]
elicbdf["Device Make"] = F["elicb_device_make"]
elicbdf["Device Type"] = F["elicb_device_type"]
elicbdf["Device Sensitivity (mA)"] = F["device_sen"]
elicbdf["Device Rating (A)"] = F["elicb_device_rating"]
elicbdf["TMS - TDS"] = F["eli_tms"]
elicbdf["Type of Circuit Location"] = F["type_cir_loc"]
elicbdf["No. of Phases"] = F["no_phases"]
elicbdf["Trip Curve"] = F["b_curve_type"]
elicbdf["V_LN (V)"] = F["elicb_mvln"]
elicbdf["V_LE (V)"] = F["elicb_mvle"]
elicbdf["V_NE (V)"] = F["elicb_mvne"]
elicbdf["L1-ELI (O) (O)"] = F["elicb_mel1"]
elicbdf["L2-ELI (O) (O)"] = F["elicb_mel2"]
elicbdf["L3-ELI (O) (O)"] = F["elicb_mel3"]
elicbdf["Psc (kA)"] = F["elicb_psc"]
elicbdf["Suggested Max ELI (O)"] = F["elicb_max_eli"]

elicbdf1 = elicbdf[
    [
        "Device Make",
        "Device Type",
        "Earthing Configuration",
        "Type of Circuit Location",
        "Device Rating (A)",
        "Device Sensitivity (mA)",
        "No. of Phases",
        "Trip Curve",
    ]
]

elicbdf2 = elicbdf[
    [
        "Device Make",
        "Device Rating (A)",
        "Device Type",
        "No. of Phases",
        "V_LN (V)",
        "V_LE (V)",
        "V_NE (V)",
        "L1-ELI (O) (O)",
        "L2-ELI (O) (O)",
        "L3-ELI (O) (O)",
        "Psc (kA)",
    ]
]


def eli_test_result1(row):
    Is = row["Device Rating (A)"]
    TMS = row["TMS - TDS"]
    Td = 0.4 if row["Type of Circuit Location"] == "Final" else 5
    K = (
        0.14
        if row["Trip Curve"] == "IEC Standard Inverse"
        else 13.5
        if row["Trip Curve"] == "IEC Very Inverse"
        else 120
        if row["Trip Curve"] == "IEC Long-Time Inverse"
        else 80
        if row["Trip Curve"] == "IEC Extremely Inverse"
        else 315.2
        if row["Trip Curve"] == "IEC Ultra Inverse"
        else None
    )
    P = (
        0.02
        if row["Trip Curve"] == "IEC Standard Inverse"
        else 1
        if row["Trip Curve"] == "IEC Very Inverse"
        else 1
        if row["Trip Curve"] == "IEC Long-Time Inverse"
        else 2
        if row["Trip Curve"] == "IEC Extremely Inverse"
        else 2.5
        if row["Trip Curve"] == "IEC Ultra Inverse"
        else None
    )

    if row["Earthing Configuration"] == "TN":
        IEC_val_TN = row["V_LE (V)"] / (Is * ((((K * TMS) / Td) + 1) ** (1 / P)))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI (O) (O)"] <= IEC_val_TN
                and row["L2-ELI (O) (O)"] <= IEC_val_TN
                and row["L3-ELI (O) (O)"] <= IEC_val_TN
            ):
                return "Pass"
            else:
                return "Fail"
        elif row["No. of Phases"] == 1:
            if row["L1-ELI (O) (O)"] <= IEC_val_TN:
                return "Pass"
            else:
                return "Fail"
        else:
            return "N/A"
    elif row["Earthing Configuration"] == "TT":
        IEC_val_TT = 50 / (Is * ((((K * TMS) / Td) + 1) ** (1 / P)))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI (O) (O)"] <= IEC_val_TT
                and row["L2-ELI (O) (O)"] <= IEC_val_TT
                and row["L3-ELI (O) (O)"] <= IEC_val_TT
            ):
                return "Pass"
            else:
                return "Fail"
        elif row["No. of Phases"] == 1:
            if row["L1-ELI (O) (O)"] <= IEC_val_TT:
                return "Pass"
            else:
                return "Fail"
        else:
            return "N/A"


def eli_test_result2(row):
    Is = row["Device Rating (A)"]
    Td = 0.4 if row["Type of Circuit Location"] == "Final" else 5
    TDS = row["TMS - TDS"]
    A = (
        0.0515
        if row["Trip Curve"] == "IEEE Moderately Inverse"
        else 19.61
        if row["Trip Curve"] == "IEEE Very Inverse"
        else 28.2
        if row["Trip Curve"] == "IEEE Extremely Inverse"
        else None
    )
    B = (
        0.114
        if row["Trip Curve"] == "IEEE Moderately Inverse"
        else 0.491
        if row["Trip Curve"] == "IEEE Very Inverse"
        else 0.1217
        if row["Trip Curve"] == "IEEE Extremely Inverse"
        else None
    )
    p = (
        0.02
        if row["Trip Curve"] == "IEEE Moderately Inverse"
        else 2
        if row["Trip Curve"] == "IEEE Very Inverse"
        else 2
        if row["Trip Curve"] == "IEEE Extremely Inverse"
        else None
    )

    if row["Earthing Configuration"] == "TN":
        IEEE_val_TN = row["V_LE (V)"] / (Is * (((((A / ((Td / TDS) - B)) + 1)) ** (1 / p))))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI (O) (O)"] <= IEEE_val_TN
                and row["L2-ELI (O) (O)"] <= IEEE_val_TN
                and row["L3-ELI (O) (O)"] <= IEEE_val_TN
            ):
                return "Pass"
            else:
                return "Fail"
        elif row["No. of Phases"] == 1:
            if row["L1-ELI (O) (O)"] <= IEEE_val_TN:
                return "Pass"
            else:
                return "Fail"
        else:
            return "N/A"
    elif row["Earthing Configuration"] == "TT":
        IEEE_val_TT = 50 / (Is * (((((A / ((Td / TDS) - B)) + 1)) ** (1 / p))))
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI (O) (O)"] <= IEEE_val_TT
                and row["L2-ELI (O) (O)"] <= IEEE_val_TT
                and row["L3-ELI (O) (O)"] <= IEEE_val_TT
            ):
                return "Pass"
            else:
                return "Fail"
        elif row["No. of Phases"] == 1:
            if row["L1-ELI (O) (O)"] <= IEEE_val_TT:
                return "Pass"
            else:
                return "Fail"
        else:
            return "N/A"


result_column1 = []
for index, row in elicbdf.iterrows():
    if row["Device Type"] == "MCB":
        rating = row["Device Rating (A)"]
        trip = row["Trip Curve"]
        result_row = fg[fg["Device Rating (A)"] == rating]
        if trip in result_row.columns:
            val_MCB = result_row[trip].values[0]
        else:
            val_MCB = (
                0  # Set a default value when 'Trip Curve' value is not found in sugg-max-eli.csv
            )
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI (O) (O)"] <= val_MCB
                and row["L2-ELI (O) (O)"] <= val_MCB
                and row["L3-ELI (O) (O)"] <= val_MCB
            ):
                result_column1.append("Pass")
            else:
                result_column1.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI (O) (O)"] <= val_MCB:
                result_column1.append("Pass")
            else:
                result_column1.append("Fail")
        else:
            result_column1.append("N/A")

    elif row["Device Type"] in ["RCD", "RCBO", "RCCB"] and row["Earthing Configuration"] == "TN":
        rccb_val_TN = (row["V_LE (V)"] / row["Device Sensitivity (mA)"]) * 1000
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI (O) (O)"] <= rccb_val_TN
                and row["L2-ELI (O) (O)"] <= rccb_val_TN
                and row["L3-ELI (O) (O)"] <= rccb_val_TN
            ):
                result_column1.append("Pass")
            else:
                result_column1.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI (O) (O)"] <= rccb_val_TN:
                result_column1.append("Pass")
            else:
                result_column1.append("Fail")
        else:
            result_column1.append("N/A")

    elif row["Device Type"] in ["RCD", "RCBO", "RCCB"] and row["Earthing Configuration"] == "TT":
        rccb_val_TT = (50 / row["Device Sensitivity (mA)"]) * 1000
        if row["No. of Phases"] == 3:
            if (
                row["L1-ELI (O) (O)"] <= rccb_val_TT
                and row["L2-ELI (O) (O)"] <= rccb_val_TT
                and row["L3-ELI (O) (O)"] <= rccb_val_TT
            ):
                result_column1.append("Pass")
            else:
                result_column1.append("Fail")
        elif row["No. of Phases"] == 1:
            if row["L1-ELI (O) (O)"] <= rccb_val_TT:
                result_column1.append("Pass")
            else:
                result_column1.append("Fail")
        else:
            result_column1.append("N/A")

elicbdf2["Result"] = result_column1


def eli_test_table1(df, doc):
    doc.add_heading("Earth Loop Impedance Test - Circuit Breaker", level=1)
    table_data = df.iloc[:, :8]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.2,
        1: 0.5,
        2: 0.6,
        3: 0.6,
        4: 0.8,
        5: 0.7,
        6: 0.6,
        7: 0.6,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for j, col in enumerate(table_data.columns):
        for i in range(num_rows):
            cell = table.cell(i + 1, j)
            cell.text = str(table_data[col][i])
            cell.width = Inches(column_widths.get(j, 1))
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def eli_test_table2(df, doc):
    doc.add_heading("Earth Loop Impedance Test - RCD, RCBO, RCCB", level=1)
    table_data = df.iloc[:, [0, 3, 2, 6, 7, 8, 9, 10, 11, 12, 13]]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False

    column_widths = {
        0: 0.5,
        1: 0.6,
        2: 0.6,
        3: 0.6,
        4: 0.6,
        5: 0.6,
        6: 0.6,
        7: 0.6,
        8: 0.6,
        9: 0.6,
        10: 0.6,
    }

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for j, col in enumerate(table_data.columns):
        for i in range(num_rows):
            cell = table.cell(i + 1, j)
            cell.text = str(table_data[col][i])
            cell.width = Inches(column_widths.get(j, 1))
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_custom_heading(heading_text, doc):
    heading = doc.add_paragraph()
    run = heading.add_run(heading_text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def add_custom_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(parse_xml('<w:br w:type="page"/>'))


def generate_report():
    doc = Document()
    doc.add_heading("Electrical Installation Test Report", level=0)

    add_custom_heading("1. Earth Loop Impedance Test Results", doc)

    eli_test_table1(elicbdf1, doc)
    add_custom_page_break(doc)
    eli_test_table2(elicbdf2, doc)

    doc.save("test_report.docx")


generate_report()
