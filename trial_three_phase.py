import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from tabulate import tabulate
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import csv
import io

F = pd.read_csv("main.csv")

tpsdf = pd.DataFrame()  # TPS
tpsdf["Rated Line Voltage (V)"] = F["rated_line_vlt"]
tpsdf["Voltage-L1L2 (V)"] = F["tps_l1_l2_v"]
tpsdf["Voltage-L2L3 (V)"] = F["tps_l2_l3_v"]
tpsdf["Voltage-L3L1 (V)"] = F["tps_l3_l1_v"]
tpsdf["Voltage-L1N (V)"] = F["tps_l1_n_v"]
tpsdf["Voltage-L2N (V)"] = F["tps_l2_n_v"]
tpsdf["Voltage-L3N (V)"] = F["tps_l3_n_v"]
tpsdf["Average Line Voltage (V)"] = F["avg_line_vlt"]
tpsdf["Average Phase Voltage (V)"] = F["avg_ph_vlt"]
tpsdf["Voltage Unbalance %"] = F["vlt_unb"]
# tpsdf["Voltage Result"] = F["tps_vlt_result"]
tpsdf["Rated Phase Current (A)"] = F["rated_ph_curr"]
tpsdf["Current-L1 (A)"] = F["tps_curr_l1"]
tpsdf["Current-L2 (A)"] = F["tps_curr_l2"]
tpsdf["Current-L3 (A)"] = F["tps_curr_l3"]
tpsdf["Average Phase Current (A)"] = F["avg_ph_curr"]
tpsdf["Current Unbalance %"] = F["curr_unb"]
# tpsdf["Current Result"] = F["tps_curr_result"]
tpsdf["Voltage-NE (V)"] = F["vlt_nev"]
# tpsdf["NEV Result"] = F["tps_nev_result"]
tpsdf["Zero Sum Current (mA)"] = F["zero_sum_curr"]
# tpsdf["ZeroSum Result"] = F["tps_zs_result"]
# tpsdf["Remarks"] = F["tps_result_remarks"]


def find_loc_name(loc_id, locations):
    while loc_id in locations:
        loc_data = locations[loc_id]
        loc_name = loc_data["loc_name"]
        granularity = int(loc_data["granularity"])
        loc_type = loc_data["type"]

        if granularity == 1:
            if loc_type != "fac_area":
                parent_id = loc_data["parent_id"]
                parent_loc_name = find_loc_name(parent_id, locations)
                return parent_loc_name if parent_loc_name else "Parent Location"
            else:
                return loc_name
        else:
            loc_id = loc_data["parent_id"]
    return None


def create_dataframe(F):
    with open("your_file.csv", "r") as file:
        reader = csv.DictReader(file)
        locations = {row["loc_id"]: row for row in reader}
    table_data = []
    for index, row in F.iterrows():
        loc_id = str(row["loc_id"])
        result = find_loc_name(loc_id, locations)
        if result:
            loc_name = locations[loc_id]["loc_name"]
            parent_id = locations[loc_id]["parent_id"]
            parent_loc_name = locations[parent_id]["loc_name"]
            table_data.append([loc_name, parent_loc_name if parent_loc_name else "", result])
    df = pd.DataFrame(table_data, columns=["Location", "Parent Location", "Facility Area"])
    return df


table_df = create_dataframe(F)


def threephase_result(tpsdf, tf2):
    tpsdf["Rated Line Voltage (V)"] = tf2["Rated Line Voltage (V)"]
    tpsdf["Average Line Voltage (V)"] = round(
        (tf2["Voltage-L1L2 (V)"] + tf2["Voltage-L2L3 (V)"] + tf2["Voltage-L3L1 (V)"]) / 3, 2
    )

    tpsdf["Average Phase Voltage (V)"] = (
        tf2["Voltage-L1N (V)"] + tf2["Voltage-L2N (V)"] + tf2["Voltage-L3N (V)"]
    ) / 3
    tpsdf["Voltage Unbalance %"] = round(
        (
            (tf2["Voltage-L1N (V)"] - tpsdf["Average Line Voltage (V)"])
            .abs()
            .where(
                (tf2["Voltage-L1N (V)"] - tpsdf["Average Line Voltage (V)"]) > 0,
                (tpsdf["Average Line Voltage (V)"] - tf2["Voltage-L1N (V)"]).abs(),
            )
            .max(axis=0)
            / tpsdf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    tpsdf["Voltage Result"] = np.where(tpsdf["Voltage Unbalance %"] <= 10, "Pass", "Fail")
    tpsdf["Rated Phase Current (A)"] = tf2["Rated Phase Current (A)"]
    tpsdf["Average Phase Current (A)"] = round(
        (tf2["Current-L1 (A)"] + tf2["Current-L2 (A)"] + tf2["Current-L3 (A)"]) / 3, 2
    )
    tpsdf["Current Unbalance %"] = round(
        (
            (tf2["Current-L1 (A)"] - tpsdf["Average Phase Current (A)"])
            .abs()
            .where(
                (tf2["Current-L1 (A)"] - tpsdf["Average Phase Current (A)"]) > 0,
                (tpsdf["Average Phase Current (A)"] - tf2["Current-L1 (A)"]).abs(),
            )
            .max(axis=0)
            / tpsdf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    return tpsdf


tpsdf["Current Result"] = np.where(tpsdf["Current Unbalance %"] <= 10, "PASS", "FAIL")
tpsdf["Voltage-NE (V)"] = tpsdf["Voltage-NE (V)"]
tpsdf["NEV Result"] = np.where(tpsdf["Voltage-NE (V)"] <= 2, "PASS", "FAIL")
tpsdf["Zero Sum Current (mA)"] = tpsdf["Zero Sum Current (mA)"]
tpsdf["ZeroSum Result"] = np.where(
    tpsdf["Zero Sum Current (mA)"] <= (tpsdf["Rated Phase Current (A)"] / 5000 * 1000),
    "PASS",
    "FAIL",
)
tpsdf = pd.concat([table_df, tpsdf], axis=1)
tpsdf = tpsdf.dropna()
print(tpsdf)

tpsdf_melted = tpsdf.melt(value_vars=tpsdf.columns, var_name="Parameter", value_name="Value")


def PF_table(df, doc):  # creates the insulation table with  result coloumn
    table_data = df. iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)  # Add +1 for the "Result" column
    table.style = "Table Grid"
    table.autofit = False
    max_word_lengths = [0] * num_cols
    for j, col in enumerate(table_data.columns):
        if not table_data[col].empty:
            max_word_length = max(table_data[col].astype(str).apply(len))
        else:
            max_word_length = 0
        max_word_lengths[j] = max_word_length

    conversion_factor = -0.2
    column_widths = [length * conversion_factor for length in max_word_lengths]

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=0):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i + 1, j).text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[
                0
            ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value in {"Pass", "PASS", "CLOCKWISE"}:
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value in {"Fail", "FAIL", "ANTICLOCKWISE"}:
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def main():
    doc = Document()
    doc.add_heading("Three Phase Symmetry test", 0)
    doc = PF_table(tpsdf, doc)
    doc.save("three.docx")


main()
