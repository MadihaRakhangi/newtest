import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from tabulate import tabulate
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import csv
import io

F = pd.read_csv("main.csv")


# --------------------------------------------location------------------
def find_loc_name(loc_id, locations):
    while loc_id in locations:
        loc_data = locations[loc_id]
        loc_name = loc_data["loc_name"]
        granularity = int(loc_data["granularity"])
        loc_type = loc_data["type"]

        if granularity == 1:
            if loc_type != "fac_area":
                parent_id = loc_data["parent_id"]
                parent_loc_name = find_loc_name(parent_id, locations)
                return parent_loc_name if parent_loc_name else "Parent Location"
            else:
                return loc_name
        else:
            loc_id = loc_data["parent_id"]
    return None


def create_dataframe(F):
    with open("your_file.csv", "r") as file:
        reader = csv.DictReader(file)
        locations = {row["loc_id"]: row for row in reader}
    table_data = []
    for index, row in F.iterrows():
        loc_id = str(row["loc_id"])
        result = find_loc_name(loc_id, locations)
        if result:
            loc_name = locations[loc_id]["loc_name"]
            parent_id = locations[loc_id]["parent_id"]
            parent_loc_name = locations[parent_id]["loc_name"]
            table_data.append([loc_name, parent_loc_name if parent_loc_name else "", result])
    df = pd.DataFrame(table_data, columns=["Location", "Parent Location", "Facility Area"])
    return df


table_df = create_dataframe(F)


# ------------------------------------table---------------------------------------
def table(df, doc):
    table_data = df.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False
    max_word_lengths = [0] * num_cols
    for j, col in enumerate(table_data.columns):
        if not table_data[col].empty:
            max_word_length = max(table_data[col].astype(str).apply(len))
        else:
            max_word_length = 0
        max_word_lengths[j] = max_word_length

    conversion_factor = 0.1
    column_widths = [length * conversion_factor for length in max_word_lengths]

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=0):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i + 1, j).text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row in doc.tables[0].rows:
        result_cell = row.cells[-1]
        result = result_cell.text.strip()
        if result in {"Unsatisfactory", "ANTICLOCKWISE", "Fail", "REVERSE"}:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FF0000"/>')
        elif result in {"Satisfactory", "CLOCKWISE", "Pass", "OK", "Ok"}:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="00FF00"/>')
        else:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
        result_cell._element.xpath(".//w:tcPr")[0].append(shading)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 8
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


# ----------------------------------------------------Insulation resistance-----------------------
irdf = pd.DataFrame()  # IR
irdf["No. Poles"] = F["no_poles"]
irdf["SPD Applicable"] = F["spd_applicable"]
irdf["Nominal Circuit Voltage (V)"] = F["nom_cir_vlt"]
irdf["Measurement Terminals"] = F["meas_term"]
irdf["Test Voltage (V)"] = F["test_vlt"]
irdf["Conductor Type"] = F["cond_type"]
irdf["Conductor Size (sq. mm)"] = F["cond_size"]
irdf["Number of Runs"] = F["no_runs"]
irdf["Insulation Type"] = F["ins_type"]
irdf["Leakage Capacitance (nF)"] = F["leak_cap_nf"]
irdf["Insulation Resistance (MO)"] = F["ins_res_mohm"]


def insualtion_result(nom_cir_vlt, test_vlt, ins_res_mohm):
    if nom_cir_vlt <= 50:
        if ins_res_mohm >= 0.5 and test_vlt == 250:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif 50 < nom_cir_vlt <= 500:
        if ins_res_mohm >= 1 and test_vlt == 500:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif nom_cir_vlt > 500:
        if ins_res_mohm >= 1 and test_vlt == 1000:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    else:
        return "Invalid input"


irdf["Result"] = irdf.apply(
    lambda row: insualtion_result(
        row["Nominal Circuit Voltage (V)"],
        row["Test Voltage (V)"],
        row["Insulation Resistance (MO)"],
    ),
    axis=1,
)

irdf = pd.concat([table_df, irdf], axis=1)
irdf = irdf.dropna()
print(irdf)


def insulation_combined_graph(mf):
    mf = pd.read_csv("Insulate.csv")

    # Bar graph
    x = mf["Location"]
    y = mf["Nominal Circuit Voltage"]

    fig = plt.figure(figsize=(12, 6))  # Adjust the figsize as desired
    ax1 = fig.add_subplot(121)
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]  # Add more colors if needed
    ax1.bar(x, y, color=colors)
    ax1.set_xlabel("Location")
    ax1.set_ylabel("Nominal Circuit Voltage")
    ax1.set_title("Nominal Circuit Voltage by Location")

    # Pie chart
    earthing_system_counts = mf["Earthing System"].value_counts()
    ax2 = fig.add_subplot(122)
    colors = ["#5ac85a", "#dc0000"]
    ax2.pie(
        earthing_system_counts,
        labels=earthing_system_counts.index,
        autopct="%1.1f%%",
        colors=colors,
    )
    ax2.set_title("Earthing System Distribution")
    ax2.axis("equal")

    graph_combined2 = io.BytesIO()
    plt.savefig(graph_combined2)
    plt.close()

    return graph_combined2


# ------------------------------------Floor wall resistance------------------------------
fwrdf = pd.DataFrame()  # FWR
fwrdf["Distance from previous test location (m)"] = F["distance_prev_loc"]
fwrdf["Nominal Voltage to Earth of System (V)"] = F["nom_vlt_e_v"]
fwrdf["Applied Test Voltage (V)"] = F["app_test_vlt"]
fwrdf["Measured Output Current (mA)"] = F["meas_out_curr_ma"]
fwrdf["Resistance (kO)"] = F["res_kohm"]

fwrdf["Effective Resistance"] = (
    fwrdf["Applied Test Voltage (V)"] / fwrdf["Measured Output Current (mA)"]
)
fwrdf.to_csv("floorfinal.csv", index=False)


def resistance_result(nom_vlt_e_v, app_test_vlt, Eff_Floor, distance_prev_loc):
    if nom_vlt_e_v <= 500 and distance_prev_loc >= 1:
        if app_test_vlt == nom_vlt_e_v and Eff_Floor >= 50:
            return "Pass"
        else:
            return "Fail"
    elif nom_vlt_e_v > 500 and distance_prev_loc >= 1:
        if app_test_vlt == nom_vlt_e_v and Eff_Floor >= 100:
            return "Pass"
        else:
            return "Fail"
    elif distance_prev_loc <= 1:
        return "Fail"
    else:
        return "Invalid input"


fwrdf["Result"] = fwrdf.apply(
    lambda row: resistance_result(
        row["Nominal Voltage to Earth of System (V)"],
        row["Applied Test Voltage (V)"],
        row["Effective Resistance"],
        row["Distance from previous test location (m)"],
    ),
    axis=1,
)

fwrdf = pd.concat([table_df, fwrdf], axis=1)
fwrdf = fwrdf.dropna()
print(fwrdf)

# -----------------------------------------------------Resistance conductor----------------------------
rcdf = pd.DataFrame()  # RC
rcdf["Conductor Location - From"] = F["rc_route_ele_1"]
rcdf["Conductor Location - To"] = F["rc_route_ele_2"]
rcdf["No of runs of Conductor"] = F["no_runs"]
rcdf["Conductor Type"] = F["cond_type"]
rcdf["Conductor Size (sq. mm)"] = F["cond_size"]
rcdf["Conductor Length (m)"] = F["cond_len"]
rcdf["Conductor Temperature (°C)"] = F["cond_temp"]
rcdf["Is Continuity found?"] = F["cont_found"]
rcdf["Lead Internal Resistance (Ω)"] = F["lead_int_res"]
rcdf["Continuity Resistance (Ω)"] = F["con_res"]
rcdf["Corrected Continuity Resistance (Ω)"] = F["corr_cont_res"]
rcdf["Specific Conductor Resistance (MO/m) at 30°C"] = F["spec_cond_res"]


def alpha(cond_temp):  # Function to calculate alpha value based on conductor typ
    if cond_temp == "Al":
        return 0.0038
    elif cond_temp == "Cu":
        return 0.00429
    elif cond_temp == "GI":
        return 0.00641
    elif cond_temp == "SS":
        return 0.003
    else:
        return None


# Calculate Corrected Continuity Resistance
rcdf["Corrected Continuity Resistance (Ω)"] = (
    rcdf["Continuity Resistance (Ω)"] - rcdf["Lead Internal Resistance (Ω)"]
)

# Calculate Specific Conductor Resistance at 30°C
alpha_values = rcdf["Conductor Type"].apply(alpha)
rcdf["Specific Conductor Resistance (MΩ/m) at 30°C"] = rcdf[
    "Corrected Continuity Resistance (Ω)"
] / (1 + alpha_values * (rcdf["Conductor Temperature (°C)"] - 30))
rcdf["Specific Conductor Resistance (MΩ/m) at 30°C"] /= 1000000
rcdf["Specific Conductor Resistance (MΩ/m) at 30°C"] = rcdf[
    "Specific Conductor Resistance (MΩ/m) at 30°C"
].apply(lambda x: format(x, "E"))
rcdf.to_csv("resistance_updated.csv", index=False)


def resc_result(Conti, C_ContR):
    if C_ContR <= 1:
        if Conti == "Yes":
            return "Pass"
        elif Conti == "No":
            return "Check Again"
        else:
            return "Invalid"
    elif C_ContR > 1:
        if Conti == "Yes":
            return "Fail"
        else:
            return "Fail"


rcdf["Result"] = rcdf.apply(
    lambda row: resc_result(
        row["Is Continuity found?"], row["Corrected Continuity Resistance (Ω)"]
    ),
    axis=1,
)

rcdf = pd.concat([table_df, rcdf], axis=1)
rcdf = rcdf.dropna()
print(rcdf)


# ------------------------------------Phase Sequence------------------------------
phs = pd.DataFrame()  # OP
phs["VL1-L2 (V)"] = F["op_l1_l2_v"]
phs["VL2-L3 (V)"] = F["op_l2_l3_v"]
phs["VL3-L1 (V)"] = F["op_l3_l1_v"]
phs["VL1-N (V)"] = F["op_l1_n_v"]
phs["VL2-N (V)"] = F["op_l2_n_v"]
phs["VL3-N (V)"] = F["op_l3_n_v"]
phs["Phase Sequence"] = F["phase_seq"]


def phase_result(phase_seq):
    if phase_seq == "RYB":
        return "CLOCKWISE"
    elif phase_seq == "RBY":
        return "ANTICLOCKWISE"
    else:
        return ""


phs.loc[:, "Result"] = phs["Phase Sequence"].apply(phase_result)

phs_df = pd.concat([table_df, phs], axis=1)
phs_df = phs_df.dropna()
print(phs_df)


# -----------------------------------------------------Voltage drop----------------------------
vddf = pd.DataFrame()  # vlt_drop_v
vddf["Circuit Route 'From'"] = F["vd_route_ele_1"]
vddf["Circuit Route 'To'"] = F["vd_route_ele_2"]
vddf["Measured Voltage (V, L-N) 'From'"] = F["ele_1_l_n_v"]
vddf["Measured Voltage (V, L-N) 'To'"] = F["ele_2_l_n_v"]
vddf["Nominal Circuit Voltage (V)"] = F["nom_cir_vlt"]
vddf["Type of Installation Supply System"] = F["type_install_supply"]
vddf["Purpose of Supply"] = F["pur_supply"]
vddf["Conductor Type"] = F["cond_type"]
vddf["Insulation Type"] = F["ins_type"]
vddf["Cable Length (m)"] = F["cond_len"]
vddf["Calculated Voltage Drop (V)"] = F["vlt_drop_v"]
vddf["Voltage Drop %"] = F["vlt_drop_perc"]

vddf["Calculated Voltage Drop (V)"] = F["ele_1_l_n_v"] - F["ele_2_l_n_v"]
vddf["Voltage Drop %"] = (
    vddf["Calculated Voltage Drop (V)"] / vddf["Measured Voltage (V, L-N) 'From'"]
) * 100
vddf["Voltage Drop %"] = vddf["Voltage Drop %"].round(decimals=2)
# vddf.to_csv("voltage_upd.csv", index=False)

lim1 = 3
lim2 = 5
lim3 = 6
lim4 = 8


def voltage_result(vlt_drop_v, type_install_supply, pur_supply, cond_len):
    if cond_len <= 0:
        if vlt_drop_v <= lim1:
            if type_install_supply == "Public" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= lim2:
            if type_install_supply == "Public" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= lim3:
            if type_install_supply == "Private" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= lim4:
            if type_install_supply == "Private" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"

    elif 0 < cond_len <= 100:
        if vlt_drop_v <= (lim1 + (cond_len * 0.005)):
            if type_install_supply == "Public" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim2 + (cond_len * 0.005)):
            if type_install_supply == "Public" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim3 + (cond_len * 0.005)):
            if type_install_supply == "Private" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim4 + (cond_len * 0.005)):
            if type_install_supply == "Private" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"

    elif cond_len > 100:
        if vlt_drop_v <= (lim1 + 0.5):
            if type_install_supply == "Public" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim2 + 0.5):
            if type_install_supply == "Public" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim3 + 0.5):
            if type_install_supply == "Private" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim4 + 0.5):
            if type_install_supply == "Private" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"


vddf["Result"] = vddf.apply(
    lambda row: voltage_result(
        row["Voltage Drop %"],
        row["Type of Installation Supply System"],
        row["Purpose of Supply"],
        row["Cable Length (m)"],
    ),
    axis=1,
)
vddf = pd.concat([table_df, vddf], axis=1)
vddf = vddf.dropna()
print(vddf)
# -----------------------------------------------------Polarity----------------------------
ptdf = pd.DataFrame()  # PT
ptdf["Device Type"] = F["pt_device_type"]
ptdf["Type of Supply"] = F["type_supply"]
ptdf["Line to Neutral Voltage (V)"] = F["l_n_v"]
ptdf["Polarity Reference"] = F["pol_ref"]


def polarity_result(l_n_v):
    if l_n_v == 230:
        return "OK"
    else:
        return "REVERSE"


ptdf["Result"] = ptdf["Line to Neutral Voltage (V)"].apply(polarity_result)


ptdf = pd.concat([table_df, ptdf], axis=1)
ptdf = ptdf.dropna()
print(ptdf)
# -----------------------------------------------------Residual current device----------------------------
rdcdf = pd.DataFrame()  # RD
rdcdf["type_supply of Voltage Waveform"] = F["type_supply"]
rdcdf["type_supply of Earthing System"] = F["type_earth_sys"]
rdcdf["Nominal Line to Earth Voltage (V)"] = F["nom_l_e_v"]
rdcdf["Nominal Current Rating(A)"] = F["nom_curr_rating"]
rdcdf["Rated Residual Operating Current,IΔn (mA)"] = F["rated_res_op_curr"]
rdcdf["Application type_supply"] = F["appl_type"]
rdcdf["Trip curve type_supply"] = F["trip_curve_type"]
rdcdf["No. of Poles"] = F["no_poles"]
rdcdf["Test Current (mA)"] = F["test_curr_ma"]
rdcdf["Trip Current (mA)"] = F["trip_curr_ma"]
rdcdf["Trip Time (ms)"] = F["trip_time"]
rdcdf["Device Tripped"] = F["device_trip"]


def residual_result(type_supply, test_curr_ma, rated_res_op_curr, device_trip, trip_time):
    if type_supply == "AC":
        if test_curr_ma == 0.5 * rated_res_op_curr:
            if device_trip == "No":
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 1 * rated_res_op_curr and device_trip == "Yes":
            if trip_time <= 300:
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 2 * rated_res_op_curr and device_trip == "Yes":
            if trip_time <= 150:
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 5 * rated_res_op_curr and device_trip == "Yes":
            if trip_time <= 40:
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"
    elif type_supply == "A":
        if test_curr_ma == 0.5 * rated_res_op_curr:
            if device_trip == "No":
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 1 * rated_res_op_curr and device_trip == "Yes":
            if 130 <= trip_time <= 500:
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 2 * rated_res_op_curr and device_trip == "Yes":
            if 60 <= trip_time <= 200:
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 5 * rated_res_op_curr and device_trip == "Yes":
            if 50 <= trip_time <= 150:
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"
    else:
        return "Pass"


rdcdf["Result"] = rdcdf.apply(
    lambda row: residual_result(
        row["type_supply of Voltage Waveform"],
        row["Test Current (mA)"],
        row["Rated Residual Operating Current,IΔn (mA)"],
        row["Device Tripped"],
        row["Trip Time (ms)"],
    ),
    axis=1,
)
rdcdf = pd.concat([table_df, rdcdf], axis=1)
rdcdf = rdcdf.dropna()
print(rdcdf)
# -----------------------------------------------------Earth pit----------------------------
epedf = pd.DataFrame()  # EPE
epedf["No of Parallel Electrodes"] = F["no_eecp"]
epedf["Earthing Application"] = F["ea"]
epedf["Type of Earthing"] = F["type_earthing"]
epedf["Earth Electrode Depth (m)"] = F["depth_iee"]
epedf["Nearest Electrode Distance (m)"] = F["d_nee"]
epedf["Measured Earth Resistance - Individual (O)"] = F["meri"]
epedf["Calculated Earth Resistance - Individual (O)"] = F["ceri"]
epedf["Electrode Distance Ratio"] = F["elec_dis_ratio"]


def Earth_result(elec_dis_ratio, meri):
    if meri <= 2 and elec_dis_ratio >= 1:
        return "PASS - Test Electrodes are properly placed"
    elif meri <= 2 and elec_dis_ratio < 1:
        return "PASS - Test Electrodes are not properly placed"
    elif meri > 2 and elec_dis_ratio >= 1:
        return "FAIL - Test Electrodes are properly placed"
    elif meri > 2 and elec_dis_ratio < 1:
        return "FAIL - Test Electrodes are not properly placed"
    else:
        return "Invalid"


def earthpit_result(elec_dis_ratio, meri):  # earth residual test condition
    if meri <= 2 and elec_dis_ratio >= 1:
        return "PASS"
    elif meri <= 2 and elec_dis_ratio < 1:
        return "PASS"
    elif meri > 2 and elec_dis_ratio >= 1:
        return "FAIL"
    elif meri > 2 and elec_dis_ratio < 1:
        return "FAIL"
    else:
        return "Invalid"


def earth_remark_result(elec_dis_ratio, meri):  # earth residual test condition
    if meri <= 2 and elec_dis_ratio >= 1:
        return "Test Electrodes are properly placed"
    elif meri <= 2 and elec_dis_ratio < 1:
        return "Test Electrodes are not properly placed"
    elif meri > 2 and elec_dis_ratio >= 1:
        return "Test Electrodes are properly placed"
    elif meri > 2 and elec_dis_ratio < 1:
        return "Test Electrodes are not properly placed"
    else:
        return "Invalid"


epedf["Result"] = epedf.apply(
    lambda row: earthpit_result(
        row["Electrode Distance Ratio"], row["Measured Earth Resistance - Individual (O)"]
    ),
    axis=1,
)
epedf["Remark"] = epedf.apply(
    lambda row: earth_remark_result(
        row["Electrode Distance Ratio"], row["Measured Earth Resistance - Individual (O)"]
    ),
    axis=1,
)
epedf = pd.concat([table_df, epedf], axis=1)
epedf = epedf.dropna()
print(epedf)

# -----------------------------------------------------Three phase symmetry----------------------------
tpsdf = pd.DataFrame()  # TPS
tpsdf["Rated Line Voltage (V)"] = F["rated_line_vlt"]
tpsdf["Voltage-L1L2 (V)"] = F["tps_l1_l2_v"]
tpsdf["Voltage-L2L3 (V)"] = F["tps_l2_l3_v"]
tpsdf["Voltage-L3L1 (V)"] = F["tps_l3_l1_v"]
tpsdf["Voltage-L1N (V)"] = F["tps_l1_n_v"]
tpsdf["Voltage-L2N (V)"] = F["tps_l2_n_v"]
tpsdf["Voltage-L3N (V)"] = F["tps_l3_n_v"]
tpsdf["Average Line Voltage (V)"] = F["avg_line_vlt"]
tpsdf["Average Phase Voltage (V)"] = F["avg_ph_vlt"]
tpsdf["Voltage Unbalance %"] = F["vlt_unb"]
# tpsdf["Voltage Result"] = F["tps_vlt_result"]
tpsdf["Rated Phase Current (A)"] = F["rated_ph_curr"]
tpsdf["Current-L1 (A)"] = F["tps_curr_l1"]
tpsdf["Current-L2 (A)"] = F["tps_curr_l2"]
tpsdf["Current-L3 (A)"] = F["tps_curr_l3"]
tpsdf["Average Phase Current (A)"] = F["avg_ph_curr"]
tpsdf["Current Unbalance %"] = F["curr_unb"]
# tpsdf["Current Result"] = F["tps_curr_result"]
tpsdf["Voltage-NE (V)"] = F["vlt_nev"]
# tpsdf["NEV Result"] = F["tps_nev_result"]
tpsdf["Zero Sum Current (mA)"] = F["zero_sum_curr"]
# tpsdf["ZeroSum Result"] = F["tps_zs_result"]
# tpsdf["Remarks"] = F["tps_result_remarks"]


def threephase_result(tpsdf, tf2):
    tpsdf["Rated Line Voltage (V)"] = tf2["Rated Line Voltage (V)"]
    tpsdf["Average Line Voltage (V)"] = round(
        (tf2["Voltage-L1L2 (V)"] + tf2["Voltage-L2L3 (V)"] + tf2["Voltage-L3L1 (V)"]) / 3, 2
    )

    tpsdf["Average Phase Voltage (V)"] = (
        tf2["Voltage-L1N (V)"] + tf2["Voltage-L2N (V)"] + tf2["Voltage-L3N (V)"]
    ) / 3
    tpsdf["Voltage Unbalance %"] = round(
        (
            (tf2["Voltage-L1N (V)"] - tpsdf["Average Line Voltage (V)"])
            .abs()
            .where(
                (tf2["Voltage-L1N (V)"] - tpsdf["Average Line Voltage (V)"]) > 0,
                (tpsdf["Average Line Voltage (V)"] - tf2["Voltage-L1N (V)"]).abs(),
            )
            .max(axis=0)
            / tpsdf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    tpsdf["Voltage Result"] = np.where(tpsdf["Voltage Unbalance %"] <= 10, "PASS", "FAIL")
    tpsdf["Rated Phase Current (A)"] = tf2["Rated Phase Current (A)"]
    tpsdf["Average Phase Current (A)"] = round(
        (tf2["Current-L1 (A)"] + tf2["Current-L2 (A)"] + tf2["Current-L3 (A)"]) / 3, 2
    )
    tpsdf["Current Unbalance %"] = round(
        (
            (tf2["Current-L1 (A)"] - tpsdf["Average Phase Current (A)"])
            .abs()
            .where(
                (tf2["Current-L1 (A)"] - tpsdf["Average Phase Current (A)"]) > 0,
                (tpsdf["Average Phase Current (A)"] - tf2["Current-L1 (A)"]).abs(),
            )
            .max(axis=0)
            / tpsdf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    return tpsdf


tpsdf["Current Result"] = np.where(tpsdf["Current Unbalance %"] <= 10, "PASS", "FAIL")
tpsdf["Voltage-NE (V)"] = tpsdf["Voltage-NE (V)"]
tpsdf["NEV Result"] = np.where(tpsdf["Voltage-NE (V)"] <= 2, "PASS", "FAIL")
tpsdf["Zero Sum Current (mA)"] = tpsdf["Zero Sum Current (mA)"]
tpsdf["ZeroSum Result"] = np.where(
    tpsdf["Zero Sum Current (mA)"] <= (tpsdf["Rated Phase Current (A)"] / 5000 * 1000),
    "PASS",
    "FAIL",
)
tpsdf = pd.concat([table_df, tpsdf], axis=1)
tpsdf = tpsdf.dropna()
print(tpsdf)

# -----------------------------------------------------Function and operations----------------------------
fodf = pd.DataFrame()  # FUNC OPS
fodf["Device type"] = F["fo_device_type"]
fodf["Functional Check"] = F["func_check"]
fodf["Interlock check"] = F["inter_check"]


def func_ops_result(func_check, inter_check):
    if func_check == "OK" and inter_check == "OK":
        return "pass"
    elif func_check == "OK" and inter_check == "Not OK":
        return "fail"
    elif func_check == "Not OK " and inter_check == "OK":
        return "fail"
    elif func_check == "Not OK" and inter_check == "OK":
        return "fail"
    else:
        return "Invalid"


# Apply the func_ops_result function to create the new 'result' column
fodf["result"] = fodf.apply(
    lambda row: func_ops_result(row["Functional Check"], row["Interlock check"]), axis=1
)
fodf = pd.concat([table_df, fodf], axis=1)
fodf = fodf.dropna()
print(fodf)
# -----------------------------------------------------ELI circuit breaker----------------------------


# -----------------------------------------------------ELI socket----------------------------


# -----------------------------------------------------PAT----------------------------
patdf = pd.DataFrame()  # PAT
patdf["Device ID"] = F["pat_device_id"]
patdf["Device Name"] = F["pat_dev_name"]
patdf["Location"] = F["pat_location"]
patdf["Voltage Rating (V)"] = F["pat_vlt_rating"]
patdf["Fuse Rating (A)"] = F["pat_fuse_rating"]
patdf["Visual Inspection"] = F["pat_vi"]
patdf["Earth Continuity (?)"] = F["pat_ec"]
patdf["Insulation Resistance (MO)"] = F["pat_ir"]
patdf["Polarity Test"] = F["pat_pt"]
patdf["Leakage (mA)"] = F["pat_leak"]
patdf["Functional Check"] = F["pat_func_check"]

print(patdf)

# def pat_combined_graph(patdf):
#     plt.figure(figsize=(16, 8))

#     # Bar graph
#     plt.subplot(121)
#     y = patdf["Earth Continuity (?)"]
#     x = patdf["Location"]
#     colors = ["#d9534f", "#5bc0de", "#aa6f73", "#428bca"]
#     plt.bar(x, y, color=colors)
#     plt.ylabel("Earth Continuity (?)")
#     plt.xlabel("Location")
#     plt.title("Location Location VS  Earth Continuity (?) ")

#     # Pie chart
#     plt.subplot(122)
#     result_counts = patdf["Overall Result"].value_counts()
#     labels = result_counts.index
#     values = result_counts.values
#     colors = ["#5ac85a", "#dc0000"]
#     plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
#     plt.title("Residual Test Results")
#     plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
#     graph_combined = io.BytesIO()
#     plt.savefig(graph_combined)
#     plt.close()

#     return graph_combined

patdf = pd.concat([table_df, patdf], axis=1)
patdf = patdf.dropna()
print(patdf)


# -----------------------------------------------------MAIN----------------------------
def main():
    doc = Document()
    # IR
    doc.add_heading("Insulation Resistance test", 0)
    doc = table(irdf, doc)
    graph_combined = insulation_combined_graph(irdf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    # FWR
    doc.add_heading("Floor wall Resistance test", 0)
    doc = table(fwrdf, doc)
    # RC
    doc.add_heading("Resistance conductor test", 0)
    doc = table(rcdf, doc)
    # PHS
    doc.add_heading("Phase Sequence test", 0)
    doc = table(phs_df, doc)
    # VD
    doc.add_heading("Voltage Drop test", 0)
    doc = table(vddf, doc)
    # PT
    doc.add_heading("Polarity test", 0)
    doc = table(ptdf, doc)
    # RCD
    doc.add_heading("Residual current device test", 0)
    doc = table(rdcdf, doc)
    # FO
    doc.add_heading("Function and operation test", 0)
    doc = table(fodf, doc)
    # PAT
    doc.add_heading("PAT test", 0)
    doc = table(patdf, doc)
    # graph_combined = pat_combined_graph(patdf)
    # doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    doc.save("newtest1.docx")


main()
