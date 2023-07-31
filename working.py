import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from tabulate import tabulate
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import csv
import io
from datetime import datetime

F = pd.read_csv("main.csv")


# --------------------------------------------df------------------
# -#  ------Insulationresistance-------

irdf = pd.DataFrame()  # IR
irdf["No. Poles"] = F["no_poles"]
irdf["SPD Applicable"] = F["spd_applicable"]
irdf["Nominal Circuit Voltage (V)"] = F["nom_cir_vlt"]
irdf["Measurement Terminals"] = F["meas_term"]
irdf["Test Voltage (V)"] = F["test_vlt"]
irdf["Conductor Type"] = F["cond_type"]
irdf["Conductor Size (sq. mm)"] = F["cond_size"]
irdf["Number of Runs"] = F["no_runs"]
irdf["Insulation Type"] = F["ins_type"]
irdf["Leakage Capacitance (nF)"] = F["leak_cap_nf"]
irdf["Insulation Resistance (MO)"] = F["ins_res_mohm"]

#  -----Floor wall resistance------

fwrdf = pd.DataFrame()  # FWR
fwrdf["Distance from previous test location (m)"] = F["distance_prev_loc"]
fwrdf["Nominal Voltage to Earth of System (V)"] = F["nom_vlt_e_v"]
fwrdf["Applied Test Voltage (V)"] = F["app_test_vlt"]
fwrdf["Measured Output Current (mA)"] = F["meas_out_curr_ma"]
fwrdf["Resistance (kO)"] = F["res_kohm"]
# fwrdf_column_widths = [1.44, 1.5, 1.75, 2.75, 3, 2.25, 2.25, 2.5, 2, 1.5,]

fwrdf["Effective Resistance"] = (
    fwrdf["Applied Test Voltage (V)"] / fwrdf["Measured Output Current (mA)"]
)
fwrdf.to_csv("floorfinal.csv", index=False)

#  -----resistance conductor-----

rcdf = pd.DataFrame()  # RC
rcdf["Conductor Location - From"] = F["rc_route_ele_1"]
rcdf["Conductor Location - To"] = F["rc_route_ele_2"]
rcdf["No of runs of Conductor"] = F["no_runs"]
rcdf["Conductor Type"] = F["cond_type"]
rcdf["Conductor Size (sq. mm)"] = F["cond_size"]
rcdf["Conductor Length (m)"] = F["cond_len"]
rcdf["Conductor Temperature (°C)"] = F["cond_temp"]
rcdf["Is Continuity found?"] = F["cont_found"]
rcdf["Lead Internal Resistance (Ω)"] = F["lead_int_res"]
rcdf["Continuity Resistance (Ω)"] = F["con_res"]
rcdf["Corrected Continuity Resistance (Ω)"] = F["corr_cont_res"]
rcdf["Specific Conductor Resistance (MO/m) at 30°C"] = F["spec_cond_res"]

def alpha(cond_temp):  # Function to calculate alpha value based on conductor typ
    if cond_temp == "Al":
        return 0.0038
    elif cond_temp == "Cu":
        return 0.00429
    elif cond_temp == "GI":
        return 0.00641
    elif cond_temp == "SS":
        return 0.003
    else:
        return None


# Calculate Corrected Continuity Resistance
rcdf["Corrected Continuity Resistance (Ω)"] = (
    rcdf["Continuity Resistance (Ω)"] - rcdf["Lead Internal Resistance (Ω)"]
)

# Calculate Specific Conductor Resistance at 30°C
alpha_values = rcdf["Conductor Type"].apply(alpha)
rcdf["Specific Conductor Resistance (MΩ/m) at 30°C"] = rcdf["Corrected Continuity Resistance (Ω)"] / (1 + alpha_values * (rcdf["Conductor Temperature (°C)"] - 30))

rcdf["Specific Conductor Resistance (MΩ/m) at 30°C"] /= 1000000

rcdf["Specific Conductor Resistance (MΩ/m) at 30°C"] = rcdf[
    "Specific Conductor Resistance (MΩ/m) at 30°C"
].apply(lambda x: format(x, "E"))

rcdf.to_csv("resistance_updated.csv", index=False)


#--------phase sequence-----------

phs = pd.DataFrame()  # OP
phs["VL1-L2 (V)"] = F["op_l1_l2_v"]
phs["VL2-L3 (V)"] = F["op_l2_l3_v"]
phs["VL3-L1 (V)"] = F["op_l3_l1_v"]
phs["VL1-N (V)"] = F["op_l1_n_v"]
phs["VL2-N (V)"] = F["op_l2_n_v"]
phs["VL3-N (V)"] = F["op_l3_n_v"]
phs["Phase Sequence"] = F["phase_seq"]


#-------------voltagedrop-----------


vddf = pd.DataFrame()  # vlt_drop_v
vddf["Circuit Route 'From'"] = F["vd_route_ele_1"]
vddf["Circuit Route 'To'"] = F["vd_route_ele_2"]
vddf["Measured Voltage (V, L-N) 'From'"] = F["ele_1_l_n_v"]
vddf["Measured Voltage (V, L-N) 'To'"] = F["ele_2_l_n_v"]
vddf["Nominal Circuit Voltage (V)"] = F["nom_cir_vlt"]
vddf["Type of Installation Supply System"] = F["type_install_supply"]
vddf["Purpose of Supply"] = F["pur_supply"]
vddf["Conductor Type"] = F["cond_type"]
vddf["Insulation Type"] = F["ins_type"]
vddf["Cable Length (m)"] = F["cond_len"]
vddf["Calculated Voltage Drop (V)"] = F["vlt_drop_v"]
vddf["Voltage Drop %"] = F["vlt_drop_perc"]

vddf["Calculated Voltage Drop (V)"] = F["ele_1_l_n_v"] - F["ele_2_l_n_v"]
vddf["Voltage Drop %"] = (
    vddf["Calculated Voltage Drop (V)"] / vddf["Measured Voltage (V, L-N) 'From'"]
) * 100
vddf["Voltage Drop %"] = vddf["Voltage Drop %"].round(decimals=2)
# vddf.to_csv("voltage_upd.csv", index=False)

lim1 = 3
lim2 = 5
lim3 = 6
lim4 = 8

#-------------polarity---------------

ptdf = pd.DataFrame()  # PT
ptdf["Device Type"] = F["pt_device_type"]
ptdf["Type of Supply"] = F["type_supply"]
ptdf["Line to Neutral Voltage (V)"] = F["l_n_v"]
ptdf["Polarity Reference"] = F["pol_ref"]

#----------residual current device---------

rdcdf = pd.DataFrame()  # RD
rdcdf["type_supply of Voltage Waveform"] = F["type_supply"]
rdcdf["type_supply of Earthing System"] = F["type_earth_sys"]
rdcdf["Nominal Line to Earth Voltage (V)"] = F["nom_l_e_v"]
rdcdf["Nominal Current Rating(A)"] = F["nom_curr_rating"]
rdcdf["Rated Residual Operating Current,IΔn (mA)"] = F["rated_res_op_curr"]
rdcdf["Application type_supply"] = F["appl_type"]
rdcdf["Trip curve type_supply"] = F["trip_curve_type"]
rdcdf["No. of Poles"] = F["no_poles"]
rdcdf["Test Current (mA)"] = F["test_curr_ma"]
rdcdf["Trip Current (mA)"] = F["trip_curr_ma"]
rdcdf["Trip Time (ms)"] = F["trip_time"]
rdcdf["Device Tripped"] = F["device_trip"]

#----------earthpit---------


epedf = pd.DataFrame()  # EPE
epedf["No of Parallel Electrodes"] = F["no_eecp"]
epedf["Earthing Application"] = F["ea"]
epedf["Type of Earthing"] = F["type_earthing"]
epedf["Earth Electrode Depth (m)"] = F["depth_iee"]
epedf["Nearest Electrode Distance (m)"] = F["d_nee"]
epedf["Measured Earth Resistance - Individual (O)"] = F["meri"]
epedf["Calculated Earth Resistance - Individual (O)"] = F["ceri"]
epedf["Electrode Distance Ratio"] = F["elec_dis_ratio"]



#--------------three phase symmetry--------------- 

tpsdf = pd.DataFrame()  # TPS
tpsdf["Rated Line Voltage (V)"] = F["rated_line_vlt"]
tpsdf["Average Line Voltage (V)"] = F["avg_line_vlt"]
tpsdf["Average Phase Voltage (V)"] = F["avg_ph_vlt"]
tpsdf["Voltage Unbalance %"] = F["vlt_unb"]
tpsdf["Rated Phase Current (A)"] = F["rated_ph_curr"]
tpsdf["Average Phase Current (A)"] = F["avg_ph_curr"]
tpsdf["Current Unbalance %"] = F["curr_unb"]
tpsdf["Voltage-NE (V)"] = F["vlt_nev"]
tpsdf["Zero Sum Current (mA)"] = F["zero_sum_curr"]



#--------------function and opertaion--------------- 

fodf = pd.DataFrame()  # FUNC OPS
fodf["Device type"] = F["fo_device_type"]
fodf["Functional Check"] = F["func_check"]
fodf["Interlock check"] = F["inter_check"]


#--------------pat--------------- 

patdf = pd.DataFrame()  # PAT
patdf["Device ID"] = F["pat_device_id"]
patdf["Device Name"] = F["pat_dev_name"]
patdf["Location"] = F["pat_location"]
patdf["Voltage Rating (V)"] = F["pat_vlt_rating"]
patdf["Fuse Rating (A)"] = F["pat_fuse_rating"]
patdf["Visual Inspection"] = F["pat_vi"]
patdf["Earth Continuity (?)"] = F["pat_ec"]
patdf["Insulation Resistance (MO)"] = F["pat_ir"]
patdf["Polarity Test"] = F["pat_pt"]
patdf["Leakage (mA)"] = F["pat_leak"]
patdf["Functional Check"] = F["pat_func_check"]
patdf["Overall Result"] = F["result_pat"]

#--------------eli-cb--------------- 
elicb = pd.DataFrame() 
elicb["Earthing Configuration"]= F["earth_config"]
elicb["device_make"] = F["elicb_device_make"]
elicb["device_type"]= F["elicb_device_type"]
elicb["device_sensitivity"]= F["device_sen"]
elicb["Device Rating (A)"]= F["elicb_device_rating"]
# TMS - TDS	eli_tms          not there
elicb["Type of Circuit Location"]=F["type_cir_loc"]
elicb["No.of Phases"]=F["no_phases"]
elicb["Trip Curve"]=F["b_curve_type"]
elicb["V_LN (V)"]=F["elicb_mvln"]
elicb["V_NE (V)"]=F["elicb_mvne"]
elicb["V_LE (V)"]=F["elicb_mvle"]
elicb["L1-ELI (O)"]=F["elicb_mel1"]
elicb["L2-ELI (O)"]=F["elicb_mel2"]
elicb["L3-ELI (O)"]=F["elicb_mel3"]
elicb["Psc (kA)"]=F["elicb_psc"]
elicb["Suggested Max ELI (O)"]=F["elicb_max_eli"]
elicb_filled = elicb.fillna("")
elicb["Device Rating (A)"] = elicb["Device Rating (A)"].astype(int)
new_column1 = []
result_column1 = []
P = 0
K = 0
TMS = 1
TDS = 1
elicb_filled = elicb.fillna("")
elicb["Device Rating (A)"] =elicb["Device Rating (A)"].astype(int)
Is = elicb["Device Rating (A)"]

gf1 = elicb[
    [
        "Earthing Configuration",
        "Type of Circuit Location",
        "Device Rating (A)",
        "device_make",
        "device_type",
        "device_sensitivity",
        "No.of Phases",
        "Trip Curve",
    ]
]

gf2 = elicb[
    [

        "V_LN (V)",
        "V_NE (V)",
        "V_LE (V)",
        "L1-ELI (O)",
        "L2-ELI (O)",
        "L1-ELI (O)",
        "Psc (kA)",
    ]
]





# --------------------------------------------location coloumn------------------
def find_loc_name(loc_id, locations):
    while loc_id in locations:
        loc_data = locations[loc_id]
        loc_name = loc_data["loc_name"]
        granularity = int(loc_data["granularity"])
        loc_type = loc_data["type"]

        if granularity == 1:
            if loc_type != "fac_area":
                parent_id = loc_data["parent_id"]
                parent_loc_name = find_loc_name(parent_id, locations)
                return parent_loc_name if parent_loc_name else "Parent Location"
            else:
                return loc_name
        else:
            loc_id = loc_data["parent_id"]
    return None


def create_dataframe(F):
    with open("your_file.csv", "r") as file:
        reader = csv.DictReader(file)
        locations = {row["loc_id"]: row for row in reader}
    table_data = []
    for index, row in F.iterrows():
        loc_id = str(row["loc_id"])
        result = find_loc_name(loc_id, locations)
        if result:
            loc_name = locations[loc_id]["loc_name"]
            parent_id = locations[loc_id]["parent_id"]
            parent_loc_name = locations[parent_id]["loc_name"]
            table_data.append([loc_name, parent_loc_name if parent_loc_name else "", result])
    df = pd.DataFrame(table_data, columns=["Location", "Parent Location", "Facility Area"])
    return df


table_df = create_dataframe(F)

# ------------------------------------result---------------------------------------

def insualtion_result(nom_cir_vlt, test_vlt, ins_res_mohm):
    if nom_cir_vlt <= 50:
        if ins_res_mohm >= 0.5 and test_vlt == 250:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif 50 < nom_cir_vlt <= 500:
        if ins_res_mohm >= 1 and test_vlt == 500:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif nom_cir_vlt > 500:
        if ins_res_mohm >= 1 and test_vlt == 1000:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    else:
        return "Invalid input"


irdf["Result"] = irdf.apply(
    lambda row: insualtion_result(
        row["Nominal Circuit Voltage (V)"],
        row["Test Voltage (V)"],
        row["Insulation Resistance (MO)"],
    ),
    axis=1,
)

irdf = pd.concat([table_df, irdf], axis=1)
irdf = irdf.dropna()
print(irdf)


def resistance_result(nom_vlt_e_v, app_test_vlt, Eff_Floor, distance_prev_loc):
    if nom_vlt_e_v <= 500 and distance_prev_loc >= 1:
        if app_test_vlt == nom_vlt_e_v and Eff_Floor >= 50:
            return "Pass"
        else:
            return "Fail"
    elif nom_vlt_e_v > 500 and distance_prev_loc >= 1:
        if app_test_vlt == nom_vlt_e_v and Eff_Floor >= 100:
            return "Pass"
        else:
            return "Fail"
    elif distance_prev_loc <= 1:
        return "Fail"
    else:
        return "Invalid input"


fwrdf["Result"] = fwrdf.apply(
    lambda row: resistance_result(
        row["Nominal Voltage to Earth of System (V)"],
        row["Applied Test Voltage (V)"],
        row["Effective Resistance"],
        row["Distance from previous test location (m)"],
    ),
    axis=1,
)

fwrdf = pd.concat([table_df, fwrdf], axis=1)
fwrdf = fwrdf.dropna()
print(fwrdf)


def resc_result(Conti, C_ContR):
    if C_ContR <= 1:
        if Conti == "Yes":
            return "Pass"
        elif Conti == "No":
            return "Check Again"
        else:
            return "Invalid"
    elif C_ContR > 1:
        if Conti == "Yes":
            return "Fail"
        else:
            return "Fail"


rcdf["Result"] = rcdf.apply(
    lambda row: resc_result(
        row["Is Continuity found?"], row["Corrected Continuity Resistance (Ω)"]
    ),
    axis=1,
)

rcdf = pd.concat([table_df, rcdf], axis=1)
rcdf = rcdf.dropna()
print(rcdf)



def phase_result(phase_seq):
    if phase_seq == "RYB":
        return "CLOCKWISE"
    elif phase_seq == "RBY":
        return "ANTICLOCKWISE"
    else:
        return ""


phs.loc[:, "Result"] = phs["Phase Sequence"].apply(phase_result)

phs_df = pd.concat([table_df, phs], axis=1)
phs_df = phs_df.dropna()
print(phs_df)


def voltage_result(vlt_drop_v, type_install_supply, pur_supply, cond_len):
    if cond_len <= 0:
        if vlt_drop_v <= lim1:
            if type_install_supply == "Public" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= lim2:
            if type_install_supply == "Public" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= lim3:
            if type_install_supply == "Private" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= lim4:
            if type_install_supply == "Private" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"

    elif 0 < cond_len <= 100:
        if vlt_drop_v <= (lim1 + (cond_len * 0.005)):
            if type_install_supply == "Public" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim2 + (cond_len * 0.005)):
            if type_install_supply == "Public" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim3 + (cond_len * 0.005)):
            if type_install_supply == "Private" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim4 + (cond_len * 0.005)):
            if type_install_supply == "Private" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"

    elif cond_len > 100:
        if vlt_drop_v <= (lim1 + 0.5):
            if type_install_supply == "Public" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim2 + 0.5):
            if type_install_supply == "Public" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim3 + 0.5):
            if type_install_supply == "Private" and pur_supply == "Lighting":
                return "Pass"
            else:
                return "Fail"
        elif vlt_drop_v <= (lim4 + 0.5):
            if type_install_supply == "Private" and pur_supply == "Other":
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"


vddf["Result"] = vddf.apply(
    lambda row: voltage_result(
        row["Voltage Drop %"],
        row["Type of Installation Supply System"],
        row["Purpose of Supply"],
        row["Cable Length (m)"],
    ),
    axis=1,
)
vddf = pd.concat([table_df, vddf], axis=1)
vddf = vddf.dropna()
print(vddf)


def polarity_result(l_n_v):
    if l_n_v == 230:
        return "OK"
    else:
        return "REVERSE"


ptdf["Result"] = ptdf["Line to Neutral Voltage (V)"].apply(polarity_result)


ptdf = pd.concat([table_df, ptdf], axis=1)
ptdf = ptdf.dropna()
print(ptdf)

def residual_result(type_supply, test_curr_ma, rated_res_op_curr, device_trip, trip_time):
    if type_supply == "AC":
        if test_curr_ma == 0.5 * rated_res_op_curr:
            if device_trip == "No":
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 1 * rated_res_op_curr and device_trip == "Yes":
            if trip_time <= 300:
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 2 * rated_res_op_curr and device_trip == "Yes":
            if trip_time <= 150:
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 5 * rated_res_op_curr and device_trip == "Yes":
            if trip_time <= 40:
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"
    elif type_supply == "A":
        if test_curr_ma == 0.5 * rated_res_op_curr:
            if device_trip == "No":
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 1 * rated_res_op_curr and device_trip == "Yes":
            if 130 <= trip_time <= 500:
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 2 * rated_res_op_curr and device_trip == "Yes":
            if 60 <= trip_time <= 200:
                return "Pass"
            else:
                return "Fail"
        elif test_curr_ma == 5 * rated_res_op_curr and device_trip == "Yes":
            if 50 <= trip_time <= 150:
                return "Pass"
            else:
                return "Fail"
        else:
            return "Fail"
    else:
        return "Pass"


rdcdf["Result"] = rdcdf.apply(
    lambda row: residual_result(
        row["type_supply of Voltage Waveform"],
        row["Test Current (mA)"],
        row["Rated Residual Operating Current,IΔn (mA)"],
        row["Device Tripped"],
        row["Trip Time (ms)"],
    ),
    axis=1,
)
rdcdf = pd.concat([table_df, rdcdf], axis=1)
rdcdf = rdcdf.dropna()
print(rdcdf)

def Earth_result(elec_dis_ratio, meri):
    if meri <= 2 and elec_dis_ratio >= 1:
        return "PASS - Test Electrodes are properly placed"
    elif meri <= 2 and elec_dis_ratio < 1:
        return "PASS - Test Electrodes are not properly placed"
    elif meri > 2 and elec_dis_ratio >= 1:
        return "FAIL - Test Electrodes are properly placed"
    elif meri > 2 and elec_dis_ratio < 1:
        return "FAIL - Test Electrodes are not properly placed"
    else:
        return "Invalid"


def earthpit_result(elec_dis_ratio, meri):  # earth residual test condition
    if meri <= 2 and elec_dis_ratio >= 1:
        return "PASS"
    elif meri <= 2 and elec_dis_ratio < 1:
        return "PASS"
    elif meri > 2 and elec_dis_ratio >= 1:
        return "FAIL"
    elif meri > 2 and elec_dis_ratio < 1:
        return "FAIL"
    else:
        return "Invalid"


def earth_remark_result(elec_dis_ratio, meri):  # earth residual test condition
    if meri <= 2 and elec_dis_ratio >= 1:
        return "Test Electrodes are properly placed"
    elif meri <= 2 and elec_dis_ratio < 1:
        return "Test Electrodes are not properly placed"
    elif meri > 2 and elec_dis_ratio >= 1:
        return "Test Electrodes are properly placed"
    elif meri > 2 and elec_dis_ratio < 1:
        return "Test Electrodes are not properly placed"
    else:
        return "Invalid"


epedf["Result"] = epedf.apply(
    lambda row: earthpit_result(
        row["Electrode Distance Ratio"], row["Measured Earth Resistance - Individual (O)"]
    ),
    axis=1,
)
epedf["Remark"] = epedf.apply(
    lambda row: earth_remark_result(
        row["Electrode Distance Ratio"], row["Measured Earth Resistance - Individual (O)"]
    ),
    axis=1,
)
epedf = pd.concat([table_df, epedf], axis=1)
epedf = epedf.dropna()
print(epedf)


def threephase_result(tpsdf, tf2):
    tpsdf["Rated Line Voltage (V)"] = tf2["Rated Line Voltage (V)"]
    tpsdf["Average Line Voltage (V)"] = round(
        (tf2["Voltage-L1L2 (V)"] + tf2["Voltage-L2L3 (V)"] + tf2["Voltage-L3L1 (V)"]) / 3, 2
    )

    tpsdf["Average Phase Voltage (V)"] = (
        tf2["Voltage-L1N (V)"] + tf2["Voltage-L2N (V)"] + tf2["Voltage-L3N (V)"]
    ) / 3
    tpsdf["Voltage Unbalance %"] = round(
        (
            (tf2["Voltage-L1N (V)"] - tpsdf["Average Line Voltage (V)"])
            .abs()
            .where(
                (tf2["Voltage-L1N (V)"] - tpsdf["Average Line Voltage (V)"]) > 0,
                (tpsdf["Average Line Voltage (V)"] - tf2["Voltage-L1N (V)"]).abs(),
            )
            .max(axis=0)
            / tpsdf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    tpsdf["Voltage Result"] = np.where(tpsdf["Voltage Unbalance %"] <= 10, "Pass", "Fail")
    tpsdf["Rated Phase Current (A)"] = tf2["Rated Phase Current (A)"]
    tpsdf["Average Phase Current (A)"] = round(
        (tf2["Current-L1 (A)"] + tf2["Current-L2 (A)"] + tf2["Current-L3 (A)"]) / 3, 2
    )
    tpsdf["Current Unbalance %"] = round(
        (
            (tf2["Current-L1 (A)"] - tpsdf["Average Phase Current (A)"])
            .abs()
            .where(
                (tf2["Current-L1 (A)"] - tpsdf["Average Phase Current (A)"]) > 0,
                (tpsdf["Average Phase Current (A)"] - tf2["Current-L1 (A)"]).abs(),
            )
            .max(axis=0)
            / tpsdf["Average Line Voltage (V)"]
        )
        * 100,
        2,
    )
    return tpsdf

tpsdf = pd.concat([table_df, tpsdf], axis=1)
tpsdf= tpsdf.dropna()
print(irdf)


tpsdf["Current Result"] = np.where(tpsdf["Current Unbalance %"] <= 10, "PASS", "FAIL")
tpsdf["Voltage-NE (V)"] = tpsdf["Voltage-NE (V)"]
tpsdf["NEV Result"] = np.where(tpsdf["Voltage-NE (V)"] <= 2, "PASS", "FAIL")
tpsdf["Zero Sum Current (mA)"] = tpsdf["Zero Sum Current (mA)"]
tpsdf["ZeroSum Result"] = np.where(
    tpsdf["Zero Sum Current (mA)"] <= (tpsdf["Rated Phase Current (A)"] / 5000 * 1000),
    "PASS",
    "FAIL",
)
tpsdf = pd.concat([table_df, tpsdf], axis=1)
tpsdf = tpsdf.dropna()
print(tpsdf)


def func_ops_result(func_check, inter_check):
    if func_check == "OK" and inter_check == "OK":
        return "Pass"
    elif func_check == "OK" and inter_check == "Not OK":
        return "Fail"
    elif func_check == "Not OK " and inter_check == "OK":
        return "Fail"
    elif func_check == "Not OK" and inter_check == "OK":
        return "Fail"
    else:
        return "Invalid"


# Apply the func_ops_result function to create the new 'result' column
fodf["Result"] = fodf.apply(
    lambda row: func_ops_result(row["Functional Check"], row["Interlock check"]), axis=1
)
fodf = pd.concat([table_df, fodf], axis=1)
fodf = fodf.dropna()
print(fodf)


patdf = pd.concat([table_df, patdf], axis=1)
patdf = patdf.dropna()
print(patdf)




# ------------------------------------table---------------------------------------
def satisfactory_table(df, column_widths, doc):
    table_data = df.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False
    # max_word_lengths = [0] * num_cols
    # for j, col in enumerate(table_data.columns):
    #     if not table_data[col].empty:
    #         max_word_length = max(table_data[col].astype(str).apply(len))
    #     else:
    #         max_word_length = 0
    #     max_word_lengths[j] = max_word_length

    # conversion_factor = -0.2
    # column_widths = [length * conversion_factor for length in max_word_lengths]

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=0):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i + 1, j).text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for row in doc.tables[0].rows:
        result_cell = row.cells[-1]
        result = result_cell.text.strip()
        if result in {"Satisfactory"}:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="00FF00"/>')
        elif result in {"Unsatisfactory"}:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FF0000"/>')
        else:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
        result_cell._element.xpath(".//w:tcPr")[0].append(shading)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 6
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def ok_table(df, doc):
    table_data = df.iloc[:, :]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False
    max_word_lengths = [0] * num_cols
    for j, col in enumerate(table_data.columns):
        if not table_data[col].empty:
            max_word_length = max(table_data[col].astype(str).apply(len))
        else:
            max_word_length = 0
        max_word_lengths[j] = max_word_length

    conversion_factor = -0.2
    column_widths = [length * conversion_factor for length in max_word_lengths]

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=0):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i + 1, j).text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[
                0
            ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value in {"OK"}:
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value in {"REVERSE"}:
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc


def pass_table(df, doc):  # creates the insulation table with  result coloumn
    table_data = df.iloc[:, 0:]
    num_rows, num_cols = table_data.shape
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)  # Add +1 for the "Result" column
    table.style = "Table Grid"
    table.autofit = False
    max_word_lengths = [0] * num_cols
    for j, col in enumerate(table_data.columns):
        if not table_data[col].empty:
            max_word_length = max(table_data[col].astype(str).apply(len))
        else:
            max_word_length = 0
        max_word_lengths[j] = max_word_length

    conversion_factor = -0.2
    column_widths = [length * conversion_factor for length in max_word_lengths]

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=0):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i + 1, j).text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[
                0
            ].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align cell text to the center
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value in {"Pass", "PASS", "CLOCKWISE"}:
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#00FF00"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value in {"Fail", "FAIL", "ANTICLOCKWISE"}:
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#FF0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)

    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def earthpit_table(df, doc):  # creates the earthpit table with  result coloumn
    df["Electrode Distance Ratio"] = round(
        df["Nearest Electrode Distance"] / df["Earth Electrode Depth"], 2
    )
    df["Calculated Earth Resistance - Individual (Ω)"] = (
        df["Measured Earth Resistance - Individual"] * df["No. of Parallel Electrodes"]
    )

    df["Remark"] = df.apply(
        lambda row: earth_remark_result(
            row["Electrode Distance Ratio"], row["Measured Earth Resistance - Individual"]
        ),
        axis=1,
    )
    df["Result"] = df.apply(
        lambda row: earthpit_result(
            row["Electrode Distance Ratio"], row["Measured Earth Resistance - Individual"]
        ),
        axis=1,
    )

    table_data = df.iloc[:, 0:]
    num_rows, num_cols = table_data.shape[0], table_data.shape[1]
    table = doc.add_table(rows=num_rows + 1, cols=num_cols)
    num_rows, num_cols = table_data.shape
    table.style = "Table Grid"
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    table.autofit = False
    max_word_lengths = [0] * num_cols
    for j, col in enumerate(table_data.columns):
        if not table_data[col].empty:
            max_word_length = max(table_data[col].astype(str).apply(len))
        else:
            max_word_length = 0
        max_word_lengths[j] = max_word_length

    conversion_factor = -0.2
    column_widths = [length * conversion_factor for length in max_word_lengths]

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths.get(j, 1))
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])
        table.cell(0, j).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        first_row_cells = table.rows[0].cells
        for cell in first_row_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell_elem = cell._element
            tc_pr = cell_elem.get_or_add_tcPr()
            shading_elem = parse_xml(f'<w:shd {nsdecls("w")} w:fill="d9ead3"/>')
            tc_pr.append(shading_elem)

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if j == num_cols - 1:  # Apply background color only to the Result column
                result_cell = cell
                if value in {"PASS"}:
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#5ac85a"/>'.format(nsdecls("w"))
                    )  # Green color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
                elif value in {"FAIL"}:
                    shading_elm = parse_xml(
                        r'<w:shd {} w:fill="#dc0000"/>'.format(nsdecls("w"))
                    )  # Red color
                    result_cell._tc.get_or_add_tcPr().append(shading_elm)
    for section in doc.sections:
        section.left_margin = Inches(0.2)

    font_size = 7
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

    return doc


# -------------------------coloumn-----------------------

irdf_column_widths = [
    0.52,
    0.43,
    0.6,
    0.37,
    0.59,
    0.59,
    0.69,
    0.49,
    0.59,
    0.59,
    0.47,
    0.51,
    0.59,
    0.59,
    0.6,
]



# ---------------------------------------combined graphs and piechart-----------------------





def insulation_combined_graph(mf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = mf["Location"]
    y = mf["Insulation Resistance (MO)"]
    colors = ["#b967ff", "#e0a899", "#fffb96", "#428bca"]  # Add more colors if needed
    sorted_indices = np.argsort(y)  # Sort the indices based on y values
    x_sorted = [x[i] for i in sorted_indices]
    y_sorted = [y[i] for i in sorted_indices]
    bars = plt.bar(x_sorted, y_sorted, color=colors)
    plt.xlabel("Location")
    plt.ylabel("Insulation Resistance (MO)")
    plt.title("Location by Insulation Resistance (MO)")
    plt.yscale("log")
    for bar, value in zip(bars, y_sorted):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value:.2f}",  # Show only the value of 'y' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    # Pie chart
    plt.subplot(122)
    # mf["Result"] = insulation_rang(mf.shape[0])
    mf_counts = mf["Result"].value_counts()
    labels = mf_counts.index.tolist()
    values = mf_counts.values.tolist()
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=colors)
    plt.title("Test Results")
    plt.axis("equal")

    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined


def flooresistance_combined_graph(df):
    plt.figure(figsize=(16, 8))

    # bar graph
    plt.subplot(121)  # Sort the DataFrame by "Location" in ascending order
    x = df["Location"]
    y = df["Effective Resistance"]
    colors = ["#b967ff", "#e0a899", "#fffb96", "#428bca"]  # Add more colors if needed
    sorted_indices = np.argsort(y)  # Sort the indices based on y values
    x_sorted = [x[i] for i in sorted_indices]
    y_sorted = [y[i] for i in sorted_indices]
    bars = plt.bar(x_sorted, y_sorted, color=colors)
    plt.xlabel("Location")
    plt.ylabel("Effective Insulation Resistance (kΩ)")
    plt.title("Location VS Effective Floor Resistance Bar Graph")
    for bar, value in zip(bars, y_sorted):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value:.2f}",  # Show only the value of 'y' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    # Pie chart
    plt.subplot(122)
    df_counts = df["Result"].value_counts()
    labels = df_counts.index.tolist()
    values = df_counts.values.tolist()
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=colors)
    plt.title("Test Results")
    plt.axis("equal")
    # Save the combined graph as bytes
    graph_combined1 = io.BytesIO()
    plt.savefig(graph_combined1)
    plt.close()

    return graph_combined1

def resc_combined_graph(jf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = jf["Conductor Type"] + ", " + jf["Conductor Length (m)"].astype(str)
    y = jf["Corrected Continuity Resistance (Ω)"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    bars = plt.bar(x, y, color=colors)
    plt.xlabel("Conductor Type and Conductor Length (m) ")
    plt.ylabel("Corrected Continuity Resistance (Ω)")
    plt.title(
        "Conductor Type and Conductor Length (m) (V) VS Corrected Continuity Resistance (Ω) "
    )
    for bar, value in zip(bars, y):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value:.2f}",  # Show only the value of 'y' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    # Pie chart
    plt.subplot(122)
    jf_counts = jf["Result"].value_counts()
    labels = jf_counts.index.tolist()
    values = jf_counts.values.tolist()
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=colors)
    plt.axis("equal")
    plt.title("Test Results")
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined


def phase_combined_graph(pf):
    plt.figure(figsize=(16, 8))
    # Bar graph
    plt.subplot(121)
    x = pf["Phase Sequence"]
    y = (
        pf["Phase Sequence"].value_counts().values
    )  # Use value_counts().values to get the count values as an array
    colors = ["#FFFF00", "#0000FF"]  # Add more colors if needed
    bars = plt.bar(x, y, color=colors)
    plt.ylabel("Count of Phase Sequence")
    plt.xlabel("Phase sequence")
    plt.title(" Phase Sequence Test")
    plt.yticks(np.arange(0, max(y) + 1, 1))
    for bar, value in zip(bars, y):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value:.2f}",  # Show only the value of 'y' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    # Pie chart
    plt.subplot(122)
    pf_counts = pf["Result"].value_counts()
    labels = pf_counts.index.tolist()
    values = pf_counts.values.tolist()
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=colors)
    plt.axis("equal")
    plt.title("Test Results")
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined

def voltage_combined_graph(vf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = vf["Conductor Type"] + ", " + vf["Cable Length (m)"].astype(str)
    y = vf["Voltage Drop %"]
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]
    bars = plt.bar(x, y, color=colors)
    plt.ylabel("Voltage Drop %")
    plt.xlabel("Conductor type and Cable Length (m)")
    plt.title("Conductor type and Cable Length (m) VS Voltage Drop %")
    for bar, value in zip(bars, y):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value:.2f}",  # Show only the value of 'y' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    

    # Pie chart
    plt.subplot(122)
    vf_counts = vf["Result"].value_counts()
    labels = vf_counts.index.tolist()
    values = vf_counts.values.tolist()
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=colors)
    plt.axis("equal")
    plt.title("Test Results")
    # Save the combined graph as bytes
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined


def polarity_combined_graph(af):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)  # Prepare the data for the graph
    result_counts = af["Result"].value_counts().sort_index()
    x = result_counts.index
    y = result_counts.values
    colors = ["#b967ff", "#e0a899", "#fffb96", "#428bca"]
    bars = plt.bar(x, y, color=colors)
    plt.xlabel("Result")
    plt.ylabel("Count")
    plt.title("Result Counts")  # Setting y-axis ticks to integers
    plt.yticks(np.arange(y.max() + 1))
    for bar, value in zip(bars, y):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value:.2f}",  # Show only the value of 'y' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    # Pie chart
    plt.subplot(122)
    af_counts = af["Result"].value_counts()
    labels = af_counts.index.tolist()
    values = af_counts.values.tolist()
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=colors)
    plt.axis("equal")
    plt.title("Polarity Results")
    # Save the combined graph as bytes
    graph_combined4 = io.BytesIO()
    plt.savefig(graph_combined4)
    plt.close()

    return graph_combined4

def residual_combined_graph(rf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    tct = rf["Trip curve type_supply"].value_counts()
    x_values = tct.index
    y_values = tct.values
    colors = ["#b967ff", "#e0a899"]  # Add more colors if needed
    bars = plt.bar(x_values, y_values, color=colors)
    plt.xlabel("Trip curve type")
    plt.ylabel("Count")
    plt.title("Residual Test Results")
    for bar, value in zip(bars, y_values):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value:.2f}",  # Show only the value of 'y' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    # Pie chart
    plt.subplot(122)
    result_counts = rf["Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined


def Earth_combined_graph(ef):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    x = ef["Location"]
    y = ef["Measured Earth Resistance - Individual (O)"]
    colors = ["#b967ff", "#e0a899", "#fffb96", "#428bca"]  # Add more colors if needed
    sorted_indices = np.argsort(y)  # Sort the indices based on y values
    x_sorted = [x[i] for i in sorted_indices]
    y_sorted = [y[i] for i in sorted_indices]
    bars = plt.bar(x_sorted, y_sorted, color=colors)
    plt.xlabel("Location")
    plt.ylabel("Measured Earth Resistance - Individual")
    plt.title("Location VS Measured Earth Resistance - Individual graph")
    for bar, value in zip(bars, y_sorted):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value:.2f}",  # Show only the value of 'y' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    # pie diagram
    plt.subplot(122)
    result_counts = ef["Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Earth Pit Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()
    return graph_combined


def func_ops_combined_graph(of):
    plt.figure(figsize=(16, 8))

    # bar graph
    plt.subplot(121)
    tct = of["Functional Check"].value_counts()
    tct1 = of["Interlock check"].value_counts()

    X = tct.index
    Ygirls = tct.values
    Zboys = tct1.values

    X_axis = np.arange(len(X))

    plt.bar(X_axis - 0.2, Ygirls, 0.4, label="ok")
    plt.bar(X_axis + 0.2, Zboys, 0.4, label="notok")

    plt.xticks(X_axis, X)
    plt.xlabel("Groups")
    plt.ylabel("Number of Students")
    plt.title("Number of Students in each group")
    plt.legend()

    # Pie chart
    plt.subplot(122)
    of_counts = of["Result"].value_counts()
    labels = of_counts.index.tolist()
    values = of_counts.values.tolist()
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90, colors=colors)
    plt.axis("equal")
    plt.title("Test Results")
    # Save the combined graph as bytes
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined



def pat_combined_graph(bf):
    plt.figure(figsize=(16, 8))

    # Bar graph
    plt.subplot(121)
    combined_data = pd.concat(
        [bf["Functional Check"], bf["Visual Inspection"]], keys=["col1", "col2"]
    )
    tct = combined_data.value_counts()
    x_values = tct.index
    y_values = tct.values
    colors = ["#b967ff", "#e0a899"]
    bars = plt.bar(x_values, y_values, color=colors)
    plt.xlabel("functional and inspection check")
    plt.ylabel("Count")
    plt.title("Residual Test Results")

    # Add text labels inside each bar
    for bar, value in zip(bars, y_values):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height / 2,
            f"{value}",  # Show the count of 'col1' inside the bar
            ha="center",
            va="center",
            color="black",
            fontsize=8,
            rotation=0,
        )

    # Pie chart
    plt.subplot(122)
    result_counts = bf["Overall Result"].value_counts()
    labels = result_counts.index
    values = result_counts.values
    colors = ["#00FF00", "#FF0000"]
    plt.pie(values, labels=labels, autopct="%1.1f%%", shadow=False, startangle=90, colors=colors)
    plt.title("Residual Test Results")
    plt.axis("equal")  # Equal aspect ratio ensures that the pie is drawn as a circle
    graph_combined = io.BytesIO()
    plt.savefig(graph_combined)
    plt.close()

    return graph_combined



# ------------------------------------Floor wall resistance------------------------------







# -----------------------------------------------------Resistance conductor----------------------------











# ------------------------------------Phase Sequence------------------------------





# -----------------------------------------------------Voltage drop----------------------------






# -----------------------------------------------------Polarity----------------------------





# -----------------------------------------------------Residual current device----------------------------




# -----------------------------------------------------Earth pit----------------------------








# -----------------------------------------------------Three phase symmetry----------------------------
# def threephase_combined_graph(tpsdf):
#     # Create a 2x2 grid of subplots
#     fig, axes = plt.subplots(nrows=2, ncols=2, figsize=(16, 8))

#     # First subplot - Bar graph for Voltage Unbalance %
#     ax1 = axes[0, 0]
#     x = tpsdf["Location"]
#     y = tpsdf["Voltage Unbalance %"]
#     colors = ["#5cb85c", "#428bca"]
#     bars = ax1.bar(x, y, color=colors)
#     ax1.set_xlabel("Location")
#     ax1.set_ylabel("Voltage Unbalance %")
#     ax1.set_title("Location VS Voltage Unbalance %")
#     for bar, value in zip(bars, y):
#         height = bar.get_height()
#         ax1.text(
#             bar.get_x() + bar.get_width() / 2,
#             height / 2,
#             f"{value:.2f}%",
#             ha="center",
#             va="center",
#             color="black",
#             fontsize=8,
#             rotation=0,
#         )

#     # Second subplot - Pie chart for Current Unbalance %
#     ax2 = axes[0, 1]
#     x = tpsdf["Location"]
#     y = tpsdf["Current Unbalance %"]
#     colors = ["#d9534f", "#5bc0de"]
#     pie_wedges, _ = ax2.pie(y, colors=colors, labels=x, autopct="%.2f%%")
#     ax2.set_xlabel("Location")
#     ax2.set_ylabel("Current Unbalance %")
#     ax2.set_title("Location VS Current Unbalance %")
#     ax2.legend(pie_wedges, x, title="Location", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

#     # Third subplot - Bar graph for Voltage-NE (V)
#     ax3 = axes[1, 0]
#     x = tpsdf["Location"]
#     y = tpsdf["Voltage-NE (V)"]
#     colors = ["#5bc0de", "#5cb85c"]
#     bars = ax3.bar(x, y, color=colors)
#     ax3.set_xlabel("Location")
#     ax3.set_ylabel("Voltage-NE (V)")
#     ax3.set_title("Location VS Voltage-NE (V)")
#     for bar, value in zip(bars, y):
#         height = bar.get_height()
#         ax3.text(
#             bar.get_x() + bar.get_width() / 2,
#             height / 2,
#             f"{value:.2f}V",
#             ha="center",
#             va="center",
#             color="black",
#             fontsize=8,
#             rotation=0,
#         )

#     # Fourth subplot - Bar graph for loading
#     ax4 = axes[1, 1]
#     x = tpsdf["Location"]
#     tpsdf["loading"] = tpsdf["Average Phase Current (A)"] / tpsdf["Rated Phase Current (A)"].round(2)
#     y = tpsdf["loading"]
#     colors = ["#5bc0de", "#5cb85c"]
#     bars = ax4.bar(x, y, color=colors)
#     ax4.set_ylabel("Loading")
#     ax4.set_xlabel("Location")
#     ax4.set_title("Location VS Loading")
#     for bar, value in zip(bars, y):
#         height = bar.get_height()
#         ax4.text(
#             bar.get_x() + bar.get_width() / 2,
#             height / 2,
#             f"{value:.2f}",
#             ha="center",
#             va="center",
#             color="black",
#             fontsize=8,
#             rotation=0,
#         )
#     plt.tight_layout()
#     graph_combined = io.BytesIO()
#     plt.savefig(graph_combined)
#     plt.close()
#     return graph_combined




# -----------------------------------------------------Function and operations----------------------------





# -----------------------------------------------------ELI circuit breaker----------------------------


# -----------------------------------------------------ELI socket----------------------------


# -----------------------------------------------------PAT----------------------------



# -----------------------------------------------------MAIN----------------------------
def main():
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(12)
    for section in doc.sections:
        section.left_margin = Inches(1)
    title = doc.add_heading("TESTING REPORT", 0)
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x6F, 0xA3, 0x15)

    section = doc.sections[0]
    header = section.header

    htable = header.add_table(
        1, 2, width=Inches(6)
    )  # Create a table with two cells for the pictures
    htable.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Configure the table properties
    htable.autofit = False

    cell1 = htable.cell(0, 0)  # Get the first cell in the table
    cell1.width = Inches(4)  # Adjust the width of the first cell

    left_header_image_path = "efficienergy-logo.jpg"  # Add the first picture to the first cell
    cell1_paragraph = cell1.paragraphs[0]
    cell1_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    cell1_run = cell1_paragraph.add_run()
    cell1_run.add_picture(left_header_image_path, width=Inches(1.5))

    # Get the second cell in the table
    cell2 = htable.cell(0, 1)
    cell2.width = Inches(3)  # Adjust the width of the second cell

    # Add the second picture to the second cell
    right_header_image_path = "secqr logo.png"  # Replace with the actual image file path
    cell2_paragraph = cell2.paragraphs[0]
    cell2_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cell2_run = cell2_paragraph.add_run()
    cell2_run.add_picture(right_header_image_path, width=Inches(1.3))

    # Add the footer with text and current date-time
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    current_time = datetime.now().strftime(
        "%d-%m-%Y  %H:%M:%S"
    )  # Format the current time as desired
    footer_paragraph.text = (
        f"{current_time}\n"
        "This Report is the Intellectual Property of M/s Efficienergi Consulting Pvt. Ltd. Plagiarism in Part or Full will be considered as theft of Intellectual property. The Information in this Report is to be treated as Confidential."
    )
    
    for run in footer_paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(7)
    # IR
    doc.add_heading("Insulation Resistance test", 0)
    doc = satisfactory_table(irdf, irdf_column_widths, doc)
    graph_combined = insulation_combined_graph(irdf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # FWR
    doc.add_heading("Floor wall Resistance test", 0)
    doc = pass_table(fwrdf, doc)
    graph_combined = flooresistance_combined_graph(fwrdf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # RC
    doc.add_heading("Resistance conductor test", 0)
    doc = pass_table(rcdf, doc)
    graph_combined = resc_combined_graph(rcdf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # PHS
    doc.add_heading("Phase Sequence test", 0)
    doc = pass_table(phs_df, doc)  # clockwise-anti
    graph_combined = phase_combined_graph(phs_df)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # VD
    doc.add_heading("Voltage Drop test", 0)
    doc = pass_table(vddf, doc)
    graph_combined = voltage_combined_graph(vddf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # PT
    doc.add_heading("Polarity test", 0)
    doc = ok_table(ptdf, doc)  # ok-reverse
    graph_combined = polarity_combined_graph(ptdf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # RCD
    doc.add_heading("Residual current device test", 0)
    doc = pass_table(rdcdf, doc)
    graph_combined = residual_combined_graph(rdcdf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # EP
    doc.add_heading("Earth Pit test", 0)
    doc = pass_table(epedf, doc)
    graph_combined = Earth_combined_graph(epedf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # TPS
    doc.add_heading("Three Phase Symmetry test", 0)
    doc = pass_table(tpsdf, doc)
    # graph_combined = threephase_combined_graph(tpsdf)
    # doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # FO
    doc.add_heading("Function and operation test", 0)
    doc = pass_table(fodf, doc)
    graph_combined = func_ops_combined_graph(fodf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    # PAT
    doc.add_heading("PAT test", 0)
    doc = pass_table(patdf, doc)
    graph_combined = pat_combined_graph(patdf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))

    doc.save("finaltest.docx")


main()
