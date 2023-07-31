import pandas as pd
import matplotlib.pyplot as plt
import docx
import csv
import io
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

mf = pd.read_csv("Fcsv.csv")


def insualtion_result(nom_cir_vlt, test_vlt, ins_res_mohm):
    if nom_cir_vlt <= 50:
        if ins_res_mohm >= 0.5 and test_vlt == 250:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif 50 < nom_cir_vlt <= 500:
        if ins_res_mohm >= 1 and test_vlt == 500:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    elif nom_cir_vlt > 500:
        if ins_res_mohm >= 1 and test_vlt == 1000:
            return "Satisfactory"
        else:
            return "Unsatisfactory"
    else:
        return "Invalid input"


def insulation_table(mf, doc):
    table_data = mf.iloc[:, 0:]
    num_rows, num_cols = table_data.shape

    max_word_lengths = [0] * num_cols
    for j, col in enumerate(table_data.columns):
        max_word_length = max(table_data[col].astype(str).apply(len))
        max_word_lengths[j] = max_word_length

    conversion_factor = 0.08
    column_widths = [length * conversion_factor for length in max_word_lengths]

    table = doc.add_table(rows=num_rows + 1, cols=num_cols + 1)
    table.style = "Table Grid"
    table.autofit = False

    # Use the calculated column widths
    for j, col in enumerate(table_data.columns):
        table.cell(0, j).text = col
        table.cell(0, j).width = Inches(column_widths[j])

    table.cell(0, num_cols).text = "Result"
    table.cell(0, num_cols).width = Inches(column_widths[num_cols - 1])

    for i, row in enumerate(table_data.itertuples(), start=1):
        for j, value in enumerate(row[1:], start=0):
            table.cell(i, j).text = str(value)

    for section in doc.sections:
        section.left_margin = Inches(0.2)
    font_size = 6.5

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

    return doc


def insulation_combined_graph(mf):
    mf = pd.read_csv("Insulate.csv")
    x = mf["Location"]
    y = mf["Nominal Circuit Voltage"]
    fig = plt.figure(figsize=(12, 6))  # Adjust the figsize as desired
    ax1 = fig.add_subplot(121)
    colors = ["#d9534f", "#5bc0de", "#5cb85c", "#428bca"]  # Add more colors if needed
    ax1.bar(x, y, color=colors)
    ax1.set_xlabel("Location")
    ax1.set_ylabel("Nominal Circuit Voltage")
    ax1.set_title("Nominal Circuit Voltage by Location")
    earthing_system_counts = mf["Earthing System"].value_counts()
    ax2 = fig.add_subplot(122)
    colors = ["#5ac85a", "#dc0000"]
    ax2.pie(
        earthing_system_counts,
        labels=earthing_system_counts.index,
        autopct="%1.1f%%",
        colors=colors,
    )
    ax2.set_title("Earthing System Distribution")
    ax2.axis("equal")
    graph_combined2 = io.BytesIO()
    plt.savefig(graph_combined2)
    plt.close()
    return graph_combined2


def main():
    mf = pd.read_csv("Fcsv.csv")
    doc = docx.Document()
    doc = insulation_table(mf, doc)

    graph_combined = insulation_combined_graph(mf)
    doc.add_picture(graph_combined, width=Inches(8), height=Inches(4))
    doc.save("outputTEST.docx")


main()
